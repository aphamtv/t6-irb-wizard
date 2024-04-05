from anthropic import Anthropic
import re
import os
from rich.console import Console
from rich.panel import Panel
from pathlib import Path
import requests
from bs4 import BeautifulSoup
import asyncio
from IRB4 import haiku_cover_letter_agent
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter
from reportlab.lib.units import inch
import docx
from docx.shared import Inches
import nltk
from nltk.corpus import stopwords
from nltk.tokenize import word_tokenize
from textblob import TextBlob
from openpyxl import Workbook
import git
from github import Github
import pandas as pd
import numpy as np
import matplotlib.pyplot as plt
import seaborn as sns
from sklearn.model_selection import train_test_split
from sklearn.linear_model import LogisticRegression
from sklearn.metrics import accuracy_score, precision_score, recall_score, f1_score
from sklearn.feature_extraction.text import TfidfVectorizer
from transformers import AutoTokenizer, AutoModelForSequenceClassification, pipeline
from sentence_transformers import SentenceTransformer
from pymed import PubMed
from wordcloud import WordCloud

from pybtex.database import BibliographyData, Entry
from pybtex.plugin import find_plugin, PluginNotFound
from pybtex.style.formatting.unsrt import Style as UnsrtStyle
from pybtex.backends import latex, html
from docx import Document
from docx.shared import Inches
from docx.enum.style import WD_STYLE_TYPE

# Set up the Anthropic API client
client = Anthropic(api_key="INSERT YOUR API KEY HERE"")

# Initialize the Rich Console
console = Console()

def create_file(directory, filename, content, file_format):
    directory_path = Path(directory)
    directory_path.mkdir(parents=True, exist_ok=True)
    filepath = directory_path / filename

    if file_format == "pdf":
        create_pdf(filepath, content)
    elif file_format == "docx":
        create_docx(filepath, content)
    elif file_format == "xlsx":
        create_xlsx(filepath, content)
    else:
        with open(filepath, 'w', encoding='utf-8') as file:
            file.write(content)

    console.print(f"[green]Created file: {filepath}[/green]")

def create_pdf(filepath, content):
    pdf = canvas.Canvas(str(filepath), pagesize=letter)
    pdf.setFont("Helvetica", 12)

    text_lines = content.split("\n")
    y = 800
    for line in text_lines:
        pdf.drawString(50, y, line)
        y -= 20
        if y < 50:
            pdf.showPage()
            y = 800

    pdf.save()

def create_docx(filepath, content):
    doc = docx.Document()
    doc.add_heading("IRB Submission Document", 0)

    for paragraph in content.split("\n\n"):
        doc.add_paragraph(paragraph)

    doc.save(str(filepath))

def create_xlsx(filepath, content):
    wb = Workbook()
    ws = wb.active
    ws.title = "IRB Submission Data"

    rows = content.split("\n")
    for row in rows:
        ws.append(row.split(","))

    wb.save(str(filepath))

def opus_orchestrator(objective, previous_results=None):
    console.print(f"\n[bold]Calling Opus for objective: {objective}[/bold]")
    previous_results_text = "\n".join(previous_results) if previous_results else "None"
    messages = [
        {
            "role": "user",
            "content": f"Based on the following objective and the previous sub-task results (if any), please break down the objective into the next sub-task, and create a concise and detailed prompt for a subagent to execute that task. The prompt should include any necessary context and specific instructions. If additional research is needed to complete the task, include 'Research:' followed by the research query. Assess if the objective has been fully achieved. If the previous sub-task results comprehensively address all aspects of the objective, include the phrase 'The task is complete:' at the beginning of your response. If the objective is not yet fully achieved, break it down into the next sub-task and create a prompt for a subagent to execute that task:\n\nObjective: {objective}\n\nPrevious sub-task results:\n{previous_results_text}"
        }
    ]

    opus_response = client.messages.create(
        model="claude-3-opus-20240229",
        max_tokens=2096,
        messages=messages
    )

    response_text = opus_response.content[0].text
    console.print(Panel(response_text, title=f"[bold green]Opus Orchestrator[/bold green]", title_align="left", border_style="green", subtitle="Sending task to Haiku 👇"))
    return response_text

def haiku_sub_agent(prompt, previous_haiku_tasks=None):
    if previous_haiku_tasks is None:
        previous_haiku_tasks = []

    system_message = "Previous Haiku tasks:\n" + "\n".join(previous_haiku_tasks)

    messages = [
        {
            "role": "user",
            "content": prompt
        }
    ]

    haiku_response = client.messages.create(
        model="claude-3-haiku-20240307",
        max_tokens=4096,
        messages=messages,
        system=system_message
    )

    response_text = haiku_response.content[0].text
    console.print(Panel(response_text, title="[bold blue]Haiku Sub-agent Result[/bold blue]", title_align="left", border_style="blue", subtitle="Task completed, sending result to Opus 👇"))
    return response_text

def opus_refine(objective, sub_task_results):
    console.print(f"\nCalling Opus to provide the refined final output for the objective: {objective}")
    messages = [
        {
            "role": "user",
            "content": f"Objective: {objective}\n\nSub-task results:\n" + "\n".join(sub_task_results) + "\n\nPlease review and refine the sub-task results into a cohesive final output in the form of a comprehensive IRB submission package tailored for the specific IRB and protocol provided. Ensure that the package adheres to the latest guidelines and regulations, such as the FDA's 21 CFR Part 11, HIPAA, and the Common Rule. Incorporate relevant research findings and best practices from recent publications and industry standards. Include the following sections:\n\n1. Cover Letter\n2. Protocol Summary\n3. Investigator Brochure\n4. Informed Consent Forms\n5. Recruitment Materials\n6. Data Management Plan\n7. Safety Monitoring Plan\n8. Study Budget\n9. CVs and Credentials\n10. IRB Application Form\n11. Literature Review\n12. References and Citations\n13. Grant Support\n14. Ethical Considerations\n15. Statistical Analysis Plan\n16. Data Monitoring Committee Charter\n17. Clinical Trial Agreement\n18. Publication Policy\n19. Conflict of Interest Disclosure\n20. Patient-Reported Outcomes Measures\n\nEach section should have detailed, IRB-compliant content based on the provided protocol and IRB requirements. Ensure each section is saved as a separate file within an 'IRB_Submission' directory. The output should specify the directory and file name for each component. Ensure the output is well-organized, follows the IRB's guidelines, and is ready for submission."
        }
    ]

    opus_response = client.messages.create(
        model="claude-3-opus-20240229",
        max_tokens=4092,
        messages=messages
    )

    response_text = opus_response.content[0].text
    console.print(Panel(response_text, title="[bold green]Final Output[/bold green]", title_align="left", border_style="green"))
    return response_text

def haiku_researcher(prompt):
    messages = [
        {
            "role": "user",
            "content": prompt
        }
    ]

    haiku_response = client.messages.create(
        model="claude-3-haiku-20240307",
        max_tokens=4096,
        messages=messages,
        search_quality_reflection=True,
        search_quality_score=5
    )

    response_text = haiku_response.content[0].text
    console.print(Panel(response_text, title="[bold magenta]Haiku Researcher Result[/bold magenta]", title_align="left", border_style="magenta"))
    return response_text

def haiku_critic(sub_task_results):
    messages = [
        {
            "role": "user",
            "content": f"Sub-task results:\n" + "\n".join(sub_task_results) + "\n\nPlease critically evaluate the sub-task results and provide feedback on the following aspects:\n\n1. Accuracy and correctness of information\n2. Completeness and adherence to guidelines\n3. Clarity and organization\n4. Potential areas for improvement\n5. Compliance with IRB templates and requirements\n6. Readability and understandability for the target audience\n7. Alignment with the latest research and industry standards\n8. Inclusion of relevant statistics and data visualizations\n9. Thoroughness of the literature review and references\n10. Overall quality and effectiveness of the IRB submission package\n\nBe thorough and detailed in your evaluation, identifying any missing information or areas that need further research or refinement to meet IRB requirements and enhance the submission's impact and value."
        }
    ]

    haiku_response = client.messages.create(
        model="claude-3-haiku-20240307",
        max_tokens=4096,
        messages=messages
    )

    response_text = haiku_response.content[0].text
    console.print(Panel(response_text, title="[bold yellow]Critic Evaluation[/bold yellow]", title_align="left", border_style="yellow"))
    return response_text

def haiku_fact_checker(sub_task_results):
    messages = [
        {
            "role": "user",
            "content": f"Sub-task results:\n" + "\n".join(sub_task_results) + "\n\nPlease fact-check the information provided in the sub-task results against authoritative sources, such as official IRB guidelines, regulatory documents, scientific literature, and the latest industry news and updates. Verify the accuracy and currency of the information, correct any inaccuracies, and provide citations for the sources used in the fact-checking process. Ensure that the IRB submission package is based on the most up-to-date and reliable information available."
        }
    ]

    haiku_response = client.messages.create(
        model="claude-3-haiku-20240307",
        max_tokens=4096,
        messages=messages
    )

    response_text = haiku_response.content[0].text
    console.print(Panel(response_text, title="[bold cyan]Fact-Checker Result[/bold cyan]", title_align="left", border_style="cyan"))
    return response_text

def haiku_literature_review_agent(prompt):
    messages = [
        {
            "role": "user",
            "content": f"Based on the provided protocol, please conduct a comprehensive literature review using relevant databases (e.g., PubMed, ClinicalTrials.gov, Scopus) to identify key studies, guidelines, regulations, and recent developments related to the study. Summarize the findings and their relevance to the IRB submission, highlighting any gaps in the current knowledge or areas for further research. Ensure that the literature review is well-organized, provides strong support for the study rationale and design, and incorporates the latest research and industry insights."
        }
    ]

    haiku_response = client.messages.create(
        model="claude-3-haiku-20240307",
        max_tokens=4096,
        messages=messages
    )

    response_text = haiku_response.content[0].text
    console.print(Panel(response_text, title="[bold purple]Literature Review[/bold purple]", title_align="left", border_style="purple"))
    return response_text

def haiku_template_compliance_agent(document, irb_template):
    messages = [
        {
            "role": "user",
            "content": f"Please review the following IRB submission document:\n\n{document}\n\nEnsure that it follows the formatting and content requirements specified in the IRB template:\n\n{irb_template}\n\nProvide feedback on any areas where the document deviates from the template and suggest necessary changes to ensure compliance. Also, check that the document adheres to the latest IRB guidelines and regulatory requirements, such as the FDA's 21 CFR Part 11, HIPAA, and the Common Rule. Recommend any updates or additions needed to meet these standards."
        }
    ]

    haiku_response = client.messages.create(
        model="claude-3-haiku-20240307",
        max_tokens=4096,
        messages=messages
    )

    response_text = haiku_response.content[0].text
    console.print(Panel(response_text, title="[bold orange]Template Compliance Review[/bold orange]", title_align="left", border_style="orange"))
    return response_text

def haiku_readability_assessment_agent(document):
    messages = [
        {
            "role": "user",
            "content": f"Please assess the readability of the following IRB submission document:\n\n{document}\n\nUse established readability metrics such as Flesch-Kincaid, Gunning Fog Index, and SMOG Index to evaluate the document's readability for the target audience. Provide suggestions for improving the language, structure, and overall clarity to ensure that it is easily understandable by study participants, IRB reviewers, and other stakeholders. Aim for a readability level that is appropriate for the general public, while still maintaining the necessary technical accuracy and completeness."
        }
    ]

    haiku_response = client.messages.create(
        model="claude-3-haiku-20240307",
        max_tokens=4096,
        messages=messages
    )

    response_text = haiku_response.content[0].text
    console.print(Panel(response_text, title="[bold teal]Readability Assessment[/bold teal]", title_align="left", border_style="teal"))
    return response_text

def haiku_grant_support_agent(prompt):
    messages = [
        {
            "role": "user",
            "content": f"Based on the provided protocol, please generate a detailed section on the funding sources and grant support for the study. Include information on the specific grants, funding agencies, and any relevant budget details. Ensure that the grant information aligns with the study budget and supports the feasibility of the research. Also, provide guidance on how to effectively communicate the study's potential impact and alignment with the funding agency's goals and priorities. Suggest strategies for securing additional funding if needed, and discuss the importance of transparent and accurate reporting of grant support in the IRB submission."
        }
    ]

    haiku_response = client.messages.create(
        model="claude-3-haiku-20240307",
        max_tokens=4096,
        messages=messages
    )

    response_text = haiku_response.content[0].text
    console.print(Panel(response_text, title="[bold red]Grant Support[/bold red]", title_align="left", border_style="red"))
    return response_text

def haiku_references_citations_agent(prompt, literature_review):
    messages = [
        {
            "role": "user",
            "content": f"Based on the provided protocol and the literature review below, please generate a comprehensive reference list and in-text citations for the IRB submission documents. Ensure that the references are formatted according to the IRB's required citation style (e.g., APA, AMA, Vancouver) and that all cited sources are included in the reference list. Use a consistent and accurate citation management tool, such as Zotero or Mendeley, to organize and format the references. Provide guidance on how to effectively integrate the citations into the submission documents to support key points and demonstrate the study's grounding in the existing literature.\n\nLiterature Review:\n{literature_review}"
        }
    ]

    haiku_response = client.messages.create(
        model="claude-3-haiku-20240307",
        max_tokens=4096,
        messages=messages
    )

    response_text = haiku_response.content[0].text
    console.print(Panel(response_text, title="[bold red]References and Citations[/bold red]", title_align="left", border_style="red"))
    return response_text

def haiku_ethical_considerations_agent(prompt):
    messages = [
        {
            "role": "user",
            "content": f"Based on the provided protocol, please identify and address the key ethical considerations related to the study. Include a thorough discussion of participant privacy and confidentiality, informed consent procedures, potential risks and benefits, and any other relevant ethical issues. Ensure that the ethical considerations section demonstrates a deep understanding of the IRB's requirements, the study's specific ethical challenges, and the latest ethical guidelines and best practices in clinical research. Provide recommendations for how to effectively communicate and manage these ethical considerations throughout the study lifecycle, from participant recruitment to data sharing and dissemination of results."
        }
    ]

    haiku_response = client.messages.create(
        model="claude-3-haiku-20240307",
        max_tokens=4096,
        messages=messages
    )

    response_text = haiku_response.content[0].text
    console.print(Panel(response_text, title="[bold red]Ethical Considerations[/bold red]", title_align="left", border_style="red"))
    return response_text

def haiku_statistical_analysis_plan_agent(prompt):
    messages = [
        {
            "role": "user",
            "content": f"Based on the provided protocol, please develop a comprehensive statistical analysis plan (SAP) for the study. The SAP should include a detailed description of the study's primary and secondary endpoints, sample size calculations, statistical methods and models to be used, handling of missing data, and any planned subgroup or sensitivity analyses. Ensure that the SAP aligns with the study's objectives, design, and data collection procedures, and that it adheres to the latest statistical guidelines and best practices for clinical trials. Provide clear justifications for the chosen statistical approaches and discuss how the results will be interpreted and reported in the context of the study's research questions and hypotheses."
        }
    ]

    haiku_response = client.messages.create(
        model="claude-3-haiku-20240307",
        max_tokens=4096,
        messages=messages
    )

    response_text = haiku_response.content[0].text
    console.print(Panel(response_text, title="[bold red]Statistical Analysis Plan[/bold red]", title_align="left", border_style="red"))
    return response_text

def haiku_data_monitoring_committee_charter_agent(prompt):
    messages = [
        {
            "role": "user",
            "content": f"Based on the provided protocol, please draft a detailed charter for the study's Data Monitoring Committee (DMC). The DMC charter should outline the committee's roles and responsibilities, membership composition, meeting frequency, and decision-making processes. Describe the DMC's primary functions, such as reviewing study progress, monitoring safety data, and making recommendations regarding study continuation, modification, or termination. Ensure that the DMC charter aligns with the study's objectives, design, and regulatory requirements, and that it incorporates best practices for independent data monitoring in clinical trials. Discuss strategies for effective communication between the DMC, study sponsors, and IRB, and provide guidance on how to maintain the DMC's independence and objectivity throughout the study."
        }
    ]

    haiku_response = client.messages.create(
        model="claude-3-haiku-20240307",
        max_tokens=4096,
        messages=messages
    )

    response_text = haiku_response.content[0].text
    console.print(Panel(response_text, title="[bold red]Data Monitoring Committee Charter[/bold red]", title_align="left", border_style="red"))
    return response_text

def haiku_clinical_trial_agreement_agent(prompt):
    messages = [
        {
            "role": "user",
            "content": f"Based on the provided protocol, please generate a comprehensive clinical trial agreement (CTA) template for the study. The CTA should outline the key terms and conditions governing the relationship between the study sponsor, participating sites, and investigators. Include sections on intellectual property rights, data ownership and sharing, publication rights, indemnification, and confidentiality. Ensure that the CTA template complies with relevant legal and regulatory requirements, such as HIPAA and the Common Rule, and that it incorporates best practices for clinical trial contracting. Provide guidance on how to customize the CTA template for the specific needs and circumstances of the study, and discuss strategies for negotiating and executing the agreement with participating sites and investigators."
        }
    ]

    haiku_response = client.messages.create(
        model="claude-3-haiku-20240307",
        max_tokens=4096,
        messages=messages
    )

    response_text = haiku_response.content[0].text
    console.print(Panel(response_text, title="[bold red]Clinical Trial Agreement[/bold red]", title_align="left", border_style="red"))
    return response_text

def haiku_protocol_summary_agent(prompt):
    messages = [
        {
            "role": "user",
            "content": f"Based on the provided protocol, please generate a concise and clear protocol summary for the IRB submission. The summary should include the study's objectives, design, methodology, study population, and any potential risks and benefits to participants. Ensure that the summary is easily understandable and adheres to IRB guidelines."
        }
    ]

    haiku_response = client.messages.create(
        model="claude-3-haiku-20240307",
        max_tokens=4096,
        messages=messages
    )

    response_text = haiku_response.content[0].text
    console.print(Panel(response_text, title="[bold yellow]Protocol Summary[/bold yellow]", title_align="left", border_style="yellow"))
    return response_text

def haiku_investigator_brochure_agent(prompt):
    messages = [
        {
            "role": "user",
            "content": f"Based on the provided protocol, please generate a detailed investigator brochure for the IRB submission. The brochure should include information about the study drug or intervention, its mechanism of action, preclinical and clinical data, and any known or potential risks and side effects. Ensure that the brochure is comprehensive and follows the format required by the IRB."
        }
    ]

    haiku_response = client.messages.create(
        model="claude-3-haiku-20240307",
        max_tokens=4096,
        messages=messages
    )

    response_text = haiku_response.content[0].text
    console.print(Panel(response_text, title="[bold red]Investigator Brochure[/bold red]", title_align="left", border_style="red"))
    return response_text

def haiku_informed_consent_agent(prompt):
    messages = [
        {
            "role": "user",
            "content": f"Based on the provided protocol and IRB requirements, please generate detailed informed consent forms for the IRB submission. The forms should include all the necessary elements, such as study purpose, procedures, risks, benefits, confidentiality, and voluntary participation. Ensure that the language is easy to understand and follows the IRB's template and guidelines."
        }
    ]

    haiku_response = client.messages.create(
        model="claude-3-haiku-20240307",
        max_tokens=4096,
        messages=messages
    )

    response_text = haiku_response.content[0].text
    console.print(Panel(response_text, title="[bold red]Informed Consent Forms[/bold red]", title_align="left", border_style="red"))
    return response_text

def haiku_recruitment_materials_agent(prompt):
    messages = [
        {
            "role": "user",
            "content": f"Based on the provided protocol and IRB requirements, please generate appropriate recruitment materials for the IRB submission. The materials may include flyers, brochures, or social media posts, depending on the study's target population and recruitment strategies. Ensure that the materials are engaging, informative, and adhere to ethical guidelines and IRB requirements."
        }
    ]

    haiku_response = client.messages.create(
        model="claude-3-haiku-20240307",
        max_tokens=4096,
        messages=messages
    )

    response_text = haiku_response.content[0].text
    console.print(Panel(response_text, title="[bold red]Recruitment Materials[/bold red]", title_align="left", border_style="red"))
    return response_text

def haiku_data_management_plan_agent(prompt):
    messages = [
        {
            "role": "user",
            "content": f"Based on the provided protocol, please generate a very comprehensive data management plan for the IRB submission. The plan should include details on data collection, storage, security, and sharing. Address issues such as data confidentiality, access control, and long-term preservation. Ensure that the plan complies with relevant regulations and IRB requirements."
        }
    ]

    haiku_response = client.messages.create(
        model="claude-3-haiku-20240307",
        max_tokens=4096,
        messages=messages
    )

    response_text = haiku_response.content[0].text
    console.print(Panel(response_text, title="[bold red]Data Management Plan[/bold red]", title_align="left", border_style="red"))
    return response_text

def haiku_safety_monitoring_plan_agent(prompt):
    messages = [
        {
            "role": "user",
            "content": f"Based on the provided protocol, please generate a very detailed safety monitoring plan for the IRB submission. The plan should include procedures for identifying, reporting, and managing adverse events and serious adverse events. Describe the roles and responsibilities of the safety monitoring team and the frequency of safety reviews. Ensure that the plan adheres to regulatory requirements and IRB guidelines."
        }
    ]

    haiku_response = client.messages.create(
        model="claude-3-haiku-20240307",
        max_tokens=4096,
        messages=messages
    )

    response_text = haiku_response.content[0].text
    console.print(Panel(response_text, title="[bold red]Safety Monitoring Plan[/bold red]", title_align="left", border_style="red"))
    return response_text

def haiku_study_budget_agent(prompt):
    messages = [
        {
            "role": "user",
            "content": f"Based on the provided protocol, please generate a very detailed study budget for the IRB submission. The budget should include all anticipated expenses, such as personnel costs, equipment, supplies, participant compensation, and any other relevant costs. Provide justifications for each budget item and ensure that the budget is realistic and aligns with the study's scope and duration."
        }
    ]

    haiku_response = client.messages.create(
        model="claude-3-haiku-20240307",
        max_tokens=4096,
        messages=messages
    )

    response_text = haiku_response.content[0].text
    console.print(Panel(response_text, title="[bold red]Study Budget[/bold red]", title_align="left", border_style="red"))
    return response_text

def haiku_cv_credentials_agent(prompt):
    messages = [
        {
            "role": "user",
            "content": f"Based on the provided protocol, please compile the CVs and credentials of all key personnel involved in the study for the IRB submission. Ensure that the CVs are up-to-date and highlight relevant qualifications, experience, and publications. Include any necessary certifications or licenses required for the study."
        }
    ]

    haiku_response = client.messages.create(
        model="claude-3-haiku-20240307",
        max_tokens=4096,
        messages=messages
    )

    response_text = haiku_response.content[0].text
    console.print(Panel(response_text, title="[bold red]CVs and Credentials[/bold red]", title_align="left", border_style="red"))
    return response_text

def haiku_irb_application_form_agent(prompt):
    messages = [
        {
            "role": "user",
            "content": f"Based on the provided protocol and IRB requirements, please complete the IRB application form for the study. Ensure that all sections of the form are filled out accurately and comprehensively, including study objectives, design, population, procedures, risks, benefits, and any other relevant information. Follow the IRB's specific instructions and guidelines for completing the form."
        }
    ]

    haiku_response = client.messages.create(
        model="claude-3-haiku-20240307",
        max_tokens=4096,
        messages=messages
    )

    response_text = haiku_response.content[0].text
    console.print(Panel(response_text, title="[bold red]IRB Application Form[/bold red]", title_align="left", border_style="red"))
    return response_text





def haiku_publication_policy_agent(prompt):
    messages = [
        {
            "role": "user",
            "content": f"Based on the provided protocol, please develop a comprehensive publication policy for the study. The publication policy should outline the principles and procedures governing the dissemination of study results, including authorship criteria, timelines for manuscript preparation and submission, and guidelines for data sharing and public access. Ensure that the publication policy aligns with the study's objectives, regulatory requirements, and ethical considerations, and that it incorporates best practices for responsible and transparent research dissemination. Discuss strategies for promoting collaboration and inclusivity among study team members, while also protecting the integrity and confidentiality of the research data. Provide guidance on how to navigate potential conflicts of interest and competing publication priorities, and emphasize the importance of timely and accurate reporting of study results to the scientific community and general public."
        }
    ]

    haiku_response = client.messages.create(
        model="claude-3-haiku-20240307",
        max_tokens=4096,
        messages=messages
    )

    response_text = haiku_response.content[0].text
    console.print(Panel(response_text, title="[bold red]Publication Policy[/bold red]", title_align="left", border_style="red"))
    return response_text

def haiku_conflict_of_interest_disclosure_agent(prompt):
    messages = [
        {
            "role": "user",
            "content": f"Based on the provided protocol, please generate a comprehensive conflict of interest (COI) disclosure form and management plan for the study. The COI disclosure should require all study team members, including investigators, staff, and consultants, to report any financial, professional, or personal interests that may pose a real or perceived conflict with their roles and responsibilities in the study. Provide clear definitions and examples of what constitutes a COI, and outline the process for reviewing and managing disclosed conflicts. Ensure that the COI disclosure and management plan adheres to the latest regulatory and ethical guidelines, such as the FDA's financial disclosure requirements and the NIH's COI policy. Discuss strategies for promoting transparency and accountability in COI reporting and management, and emphasize the importance of maintaining public trust and confidence in the research enterprise."
        }
    ]

    haiku_response = client.messages.create(
        model="claude-3-haiku-20240307",
        max_tokens=4096,
        messages=messages
    )

    response_text = haiku_response.content[0].text
    console.print(Panel(response_text, title="[bold red]Conflict of Interest Disclosure[/bold red]", title_align="left", border_style="red"))
    return response_text

def haiku_patient_reported_outcomes_measures_agent(prompt):
    messages = [
        {
            "role": "user",
            "content": f"Based on the provided protocol, please identify and recommend appropriate patient-reported outcomes (PRO) measures for the study. The PRO measures should be reliable, valid, and relevant to the study population and research questions. Provide a rationale for the selected PRO measures, and discuss how they will be administered, scored, and interpreted in the context of the study's objectives and endpoints. Ensure that the PRO measures are culturally and linguistically appropriate for the target population, and that they are feasible to implement within the study's timeline and budget. Discuss strategies for maximizing patient engagement and compliance with PRO data collection, and provide guidance on how to analyze and report PRO results in accordance with the latest standards and best practices, such as the CONSORT-PRO extension and the ISPOR PRO Good Research Practices."
        }
    ]

    haiku_response = client.messages.create(
        model="claude-3-haiku-20240307",
        max_tokens=4096,
        messages=messages
    )

    response_text = haiku_response.content[0].text
    console.print(Panel(response_text, title="[bold red]Patient-Reported Outcomes Measures[/bold red]", title_align="left", border_style="red"))
    return response_text

def haiku_final_review_agent(irb_submission_documents):
    messages = [
        {
            "role": "user",
            "content": f"Please conduct a final comprehensive review of the following IRB submission documents:\n\n{irb_submission_documents}\n\nEnsure that all documents are consistent, complete, and adhere to the highest standards of scientific rigor, ethical integrity, and regulatory compliance. Check for any remaining errors, inconsistencies, or areas that need improvement. Provide a detailed summary of your review, highlighting the key strengths and unique aspects of the submission package, as well as any potential weaknesses or limitations that should be addressed before submission. Offer specific recommendations for enhancing the overall quality, clarity, and persuasiveness of the IRB application, and discuss how the study's potential impact and value to the field can be most effectively communicated to reviewers and stakeholders."
        }
    ]

    haiku_response = client.messages.create(
        model="claude-3-haiku-20240307",
        max_tokens=4096,
        messages=messages
    )

    response_text = haiku_response.content[0].text
    console.print(Panel(response_text, title="[bold red]Final Review[/bold red]", title_align="left", border_style="red"))
    return response_text
def haiku_readability_assessment_agent(document):
    messages = [
        {
            "role": "user",
            "content": f"Please assess the readability of the following IRB submission document:\n\n{document}\n\nUse established readability metrics such as Flesch-Kincaid to evaluate the document's readability for the target audience. Provide suggestions for improving the language and structure to ensure that it is easily understandable."
        }
    ]

    haiku_response = client.messages.create(
        model="claude-3-haiku-20240307",
        max_tokens=4096,
        messages=messages
    )

    response_text = haiku_response.content[0].text
    console.print(Panel(response_text, title="[bold red]Readability Assessment[/bold red]", title_align="left", border_style="red"))
    return response_text

def web_search(query):
    url = "https://www.google.com/search"
    params = {"q": query}
    headers = {
        "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/93.0.4577.82 Safari/537.36"
    }
    response = requests.get(url, params=params, headers=headers)
    soup = BeautifulSoup(response.text, "html.parser")
    search_results = []
    for result in soup.select(".tF2Cxc"):
        title = result.select_one(".DKV0Md").text
        link = result.select_one(".yuRUbf a")["href"]
        snippet_element = result.select_one(".lEBKkf")
        snippet = snippet_element.get_text(strip=True) if snippet_element else ""
        search_results.append(f"Title: {title}\nURL: {link}\nSnippet: {snippet}\n")
    return "\n".join(search_results)

def literature_search(query):
    pubmed = PubMed(tool="MyTool", email="my@email.address")
    results = pubmed.query(query, max_results=10)
    search_results = []
    for article in results:
        title = article.title
        doi = article.doi
        abstract = article.abstract
        search_results.append(f"Title: {title}\nDOI: {doi}\nAbstract: {abstract}\n")
    return "\n".join(search_results)

def format_document(document, irb_template):
    # Load the IRB template document
    template = Document(irb_template)

    # Create a new document based on the template
    formatted_doc = Document()

    # Copy the styles from the template to the new document
    for style in template.styles:
        if style.type == WD_STYLE_TYPE.PARAGRAPH:
            formatted_doc.styles.add_style(style.name, WD_STYLE_TYPE.PARAGRAPH)
            formatted_doc.styles[style.name].base_style = style.base_style
            formatted_doc.styles[style.name].font.name = style.font.name
            formatted_doc.styles[style.name].font.size = style.font.size
            formatted_doc.styles[style.name].font.bold = style.font.bold
            formatted_doc.styles[style.name].font.italic = style.font.italic
            formatted_doc.styles[style.name].font.underline = style.font.underline
            formatted_doc.styles[style.name].font.color.rgb = style.font.color.rgb
            formatted_doc.styles[style.name].paragraph_format.alignment = style.paragraph_format.alignment
            formatted_doc.styles[style.name].paragraph_format.line_spacing = style.paragraph_format.line_spacing
            formatted_doc.styles[style.name].paragraph_format.space_before = style.paragraph_format.space_before
            formatted_doc.styles[style.name].paragraph_format.space_after = style.paragraph_format.space_after

    # Parse the document content and apply the corresponding styles
    for paragraph in document.split("\n\n"):
        if paragraph.startswith("#"):
            # Heading
            heading_level = paragraph.count("#")
            heading_text = paragraph.replace("#", "").strip()
            heading = formatted_doc.add_heading(heading_text, level=heading_level)
            heading.style = f"Heading {heading_level}"
        else:
            # Regular paragraph
            formatted_paragraph = formatted_doc.add_paragraph(paragraph)
            formatted_paragraph.style = "Normal"

    # Save the formatted document
    formatted_doc.save("formatted_document.docx")

def generate_references_citations(literature_review):
    # Create a new BibliographyData object
    bib_data = BibliographyData()

    # Parse the literature review and create bibliography entries
    for i, reference in enumerate(literature_review.split("\n"), start=1):
        entry_key = f"ref{i}"
        entry = Entry("article", fields={"title": reference})
        bib_data.add_entry(entry_key, entry)

    # Generate the formatted bibliography
    style = UnsrtStyle()
    formatter = style.format_bibliography(bib_data)
    backend = find_plugin("pybtex.backends", "latex")()
    bibliography = formatter.render(backend)

    # Generate the formatted citations
    citations = []
    for entry_key in bib_data.entries:
        citation = style.format_citation(entry_key, bib_data)
        citations.append(f"\\cite{{{entry_key}}}")

    # Save the formatted bibliography and citations to files
    with open("references.tex", "w") as bib_file:
        bib_file.write(bibliography)

    with open("citations.tex", "w") as cite_file:
        cite_file.write("\n".join(citations))
def create_word_cloud(text):
    wordcloud = WordCloud(width=800, height=800, background_color="white", stopwords=None, max_words=200, min_font_size=10).generate(text)
    plt.figure(figsize=(8, 8), facecolor=None)
    plt.imshow(wordcloud, interpolation="bilinear")
    plt.axis("off")
    plt.tight_layout(pad=0)
    plt.savefig("word_cloud.png")

def perform_sentiment_analysis(text):
    tokenizer = AutoTokenizer.from_pretrained("textattack/bert-base-uncased-imdb")
    model = AutoModelForSequenceClassification.from_pretrained("textattack/bert-base-uncased-imdb")
    nlp = pipeline("sentiment-analysis", model=model, tokenizer=tokenizer)
    result = nlp(text)[0]
    return result["label"], result["score"]

def generate_study_timeline(start_date, end_date):
    start_date = pd.to_datetime(start_date)
    end_date = pd.to_datetime(end_date)
    timeline = pd.date_range(start=start_date, end=end_date, freq="MS")
    timeline_df = pd.DataFrame({"Date": timeline})
    timeline_df["Milestone"] = ""
    return timeline_df


def main():
    objective = input("Please enter the desired IRB and paste the draft protocol: ")
    task_exchanges = []
    haiku_tasks = []

    while True:
        previous_results = [result for _, result in task_exchanges]
        opus_result = opus_orchestrator(objective, previous_results)

        if "The task is complete:" in opus_result:
            final_output = opus_result.replace("The task is complete:", "").strip()
            break
        else:
            sub_task_prompt = opus_result
            if "Research:" in sub_task_prompt:
                research_prompt = sub_task_prompt.split("Research:")[1].strip()
                web_search_results = web_search(research_prompt)
                literature_search_results = literature_search(research_prompt)
                sub_task_result = f"Web Search Results:\n{web_search_results}\n\nLiterature Search Results:\n{literature_search_results}"
            else:
                sub_task_result = haiku_sub_agent(sub_task_prompt, haiku_tasks)
            haiku_tasks.append(f"Task: {sub_task_prompt}\nResult: {sub_task_result}")
            task_exchanges.append((sub_task_prompt, sub_task_result))

    critic_evaluation = haiku_critic([result for _, result in task_exchanges])
    fact_checked_results = haiku_fact_checker([result for _, result in task_exchanges])

    cover_letter = haiku_cover_letter_agent(objective)
    protocol_summary = haiku_protocol_summary_agent(objective)
    investigator_brochure = haiku_investigator_brochure_agent(objective)
    informed_consent_forms = haiku_informed_consent_agent(objective)
    recruitment_materials = haiku_recruitment_materials_agent(objective)
    data_management_plan = haiku_data_management_plan_agent(objective)
    safety_monitoring_plan = haiku_safety_monitoring_plan_agent(objective)
    study_budget = haiku_study_budget_agent(objective)
    cv_credentials = haiku_cv_credentials_agent(objective)
    irb_application_form = haiku_irb_application_form_agent(objective)
    literature_review = haiku_literature_review_agent(objective)
    grant_support = haiku_grant_support_agent(objective)
    ethical_considerations = haiku_ethical_considerations_agent(objective)
    statistical_analysis_plan = haiku_statistical_analysis_plan_agent(objective)
    data_monitoring_committee_charter = haiku_data_monitoring_committee_charter_agent(objective)
    clinical_trial_agreement = haiku_clinical_trial_agreement_agent(objective)
    publication_policy = haiku_publication_policy_agent(objective)
    conflict_of_interest_disclosure = haiku_conflict_of_interest_disclosure_agent(objective)
    patient_reported_outcomes_measures = haiku_patient_reported_outcomes_measures_agent(objective)

    references_citations = haiku_references_citations_agent(objective, literature_review)

    create_file("IRB_Submission", "Cover_Letter.docx", cover_letter, "docx")
    create_file("IRB_Submission", "Protocol_Summary.docx", protocol_summary, "docx")
    create_file("IRB_Submission", "Investigator_Brochure.pdf", investigator_brochure, "pdf")
    create_file("IRB_Submission", "Informed_Consent_Forms.docx", informed_consent_forms, "docx")
    create_file("IRB_Submission", "Recruitment_Materials.pdf", recruitment_materials, "pdf")
    create_file("IRB_Submission", "Data_Management_Plan.docx", data_management_plan, "docx")
    create_file("IRB_Submission", "Safety_Monitoring_Plan.docx", safety_monitoring_plan, "docx")
    create_file("IRB_Submission", "Study_Budget.xlsx", study_budget, "xlsx")
    create_file("IRB_Submission", "CV_Credentials.pdf", cv_credentials, "pdf")
    create_file("IRB_Submission", "IRB_Application_Form.docx", irb_application_form, "docx")
    create_file("IRB_Submission", "Literature_Review.docx", literature_review, "docx")
    create_file("IRB_Submission", "Grant_Support.docx", grant_support, "docx")
    create_file("IRB_Submission", "Ethical_Considerations.docx", ethical_considerations, "docx")
    create_file("IRB_Submission", "Statistical_Analysis_Plan.docx", statistical_analysis_plan, "docx")
    create_file("IRB_Submission", "Data_Monitoring_Committee_Charter.docx", data_monitoring_committee_charter, "docx")
    create_file("IRB_Submission", "Clinical_Trial_Agreement.docx", clinical_trial_agreement, "docx")
    create_file("IRB_Submission", "Publication_Policy.docx", publication_policy, "docx")
    create_file("IRB_Submission", "Conflict_of_Interest_Disclosure.docx", conflict_of_interest_disclosure, "docx")
    create_file("IRB_Submission", "Patient_Reported_Outcomes_Measures.docx", patient_reported_outcomes_measures, "docx")
    create_file("IRB_Submission", "References_Citations.docx", references_citations, "docx")

    irb_template = "path/to/irb_template.docx"
    template_compliance_review = haiku_template_compliance_agent(irb_application_form, irb_template)


    irb_submission_documents = [
        cover_letter,
        protocol_summary,
        investigator_brochure,
        informed_consent_forms,
        recruitment_materials,
        data_management_plan,
        safety_monitoring_plan,
        study_budget,
        cv_credentials,
        irb_application_form,
        literature_review,
        grant_support,
        ethical_considerations,
        statistical_analysis_plan,
        data_monitoring_committee_charter,
        clinical_trial_agreement,
        publication_policy,
        conflict_of_interest_disclosure,
        patient_reported_outcomes_measures,
        references_citations,
    ]
    final_review = haiku_final_review_agent(irb_submission_documents)

    commit_changes(git_repo, "Generated IRB submission documents")

    pull_request_title = "IRB Submission Documents for Review"
    pull_request_description = "Please review the generated IRB submission documents."
    create_pull_request(git_repo, pull_request_title, pull_request_description)

    exchange_log = f"IRB: {objective.split(':')[0]}\n\nDraft Protocol:\n{objective.split(':')[1].strip()}\n\n"
    exchange_log += "=" * 40 + " Task Breakdown " + "=" * 40 + "\n\n"
    for i, (prompt, result) in enumerate(task_exchanges, start=1):
        exchange_log += f"Task {i}:\n"
        exchange_log += f"Prompt: {prompt}\n"
        exchange_log += f"Result: {result}\n\n"

    exchange_log += "=" * 40 + " Critic Evaluation " + "=" * 40 + "\n\n"
    exchange_log += critic_evaluation + "\n\n"

    exchange_log += "=" * 40 + " Fact-Checker Result " + "=" * 40 + "\n\n"
    exchange_log += fact_checked_results + "\n\n"

    exchange_log += "=" * 40 + " Template Compliance Review " + "=" * 40 + "\n\n"
    exchange_log += template_compliance_review + "\n\n"

    exchange_log += "=" * 40 + " Final Review " + "=" * 40 + "\n\n"
    exchange_log += final_review + "\n\n"

    filename = re.sub(r'\W+', '_', objective) + ".md"
    with open(filename, 'w') as file:
        file.write(exchange_log)
    print(f"\nFull exchange log saved to {filename}")

if __name__ == "__main__":
    main()